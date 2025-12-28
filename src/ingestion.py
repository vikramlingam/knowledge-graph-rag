import io
import fitz  # PyMuPDF
import pdfplumber
import docx
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

def extract_text(uploaded_file) -> List[Dict[str, Any]]:
    """
    Dispatcher function to extract text based on file type.
    """
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    if file_type == 'pdf':
        return extract_from_pdf(uploaded_file)
    elif file_type in ['docx', 'doc']:
        return extract_from_docx(uploaded_file)
    elif file_type == 'txt':
        return extract_from_txt(uploaded_file)
    else:
        logger.warning(f"Unsupported file type: {file_type}")
        return []

def extract_from_pdf(uploaded_file) -> List[Dict[str, Any]]:
    """
    Extracts text and tables from a PDF file using pdfplumber for tables and PyMuPDF for text speed/accuracy.
    """
    chunks = []
    
    # Use pdfplumber for table extraction
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract tables
                tables = page.extract_tables()
                table_text = ""
                if tables:
                    for table in tables:
                        # Convert table to markdown format for better LLM understanding
                        # Filter out None values
                        cleaned_table = [[str(cell) if cell is not None else "" for cell in row] for row in table]
                        if cleaned_table:
                            # Create markdown table
                            headers = cleaned_table[0]
                            rows = cleaned_table[1:]
                            
                            md_table = "| " + " | ".join(headers) + " |\n"
                            md_table += "| " + " | ".join(["---"] * len(headers)) + " |\n"
                            for row in rows:
                                md_table += "| " + " | ".join(row) + " |\n"
                            
                            table_text += f"\n\n**Table on Page {page_num + 1}:**\n{md_table}\n"

                # Extract regular text using PyMuPDF (faster and often better for layout)
                # We re-open with fitz because pdfplumber object is different
                # Ideally we could just use pdfplumber for everything but fitz is robust.
                # For simplicity in this POC, let's stick to pdfplumber for text too to avoid re-reading stream issues
                # or just use pdfplumber text extraction which is also good.
                
                text = page.extract_text()
                
                # Combine text and table text
                full_page_text = text + table_text
                
                if full_page_text.strip():
                    chunks.append({
                        "source": uploaded_file.name,
                        "page": page_num + 1,
                        "text": full_page_text
                    })
                    
    except Exception as e:
        logger.error(f"Error extracting from PDF {uploaded_file.name}: {e}")
        
    return chunks

def extract_from_docx(uploaded_file) -> List[Dict[str, Any]]:
    """
    Extracts text from a Word document.
    """
    chunks = []
    try:
        doc = docx.Document(uploaded_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Also extract tables from docx
        for table in doc.tables:
            # Simple table extraction for docx
            rows = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                rows.append(row_text)
            
            if rows:
                # Convert to markdown
                headers = rows[0]
                data_rows = rows[1:]
                md_table = "\n| " + " | ".join(headers) + " |\n"
                md_table += "| " + " | ".join(["---"] * len(headers)) + " |\n"
                for r in data_rows:
                    md_table += "| " + " | ".join(r) + " |\n"
                full_text.append(md_table)

        text = "\n".join(full_text)
        if text.strip():
             chunks.append({
                "source": uploaded_file.name,
                "page": 1, # Docx doesn't have pages in the same way, treat as one block or split by length later
                "text": text
            })
            
    except Exception as e:
        logger.error(f"Error extracting from DOCX {uploaded_file.name}: {e}")
        
    return chunks

def extract_from_txt(uploaded_file) -> List[Dict[str, Any]]:
    """
    Extracts text from a plain text file.
    """
    chunks = []
    try:
        stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
        text = stringio.read()
        if text.strip():
            chunks.append({
                "source": uploaded_file.name,
                "page": 1,
                "text": text
            })
    except Exception as e:
        logger.error(f"Error extracting from TXT {uploaded_file.name}: {e}")
        
    return chunks
